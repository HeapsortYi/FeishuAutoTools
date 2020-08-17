# coding: utf-8

from datetime import date
import pandas as pd
import os.path as osp
from docx import Document
import sys

sys.path.append(osp.dirname(osp.dirname(__file__)))

from biz import conf as cfg
from biz.log_conf import logger


def _is_chinese(uchar):
    """判断一个unicode是否是汉字"""
    return '\u4e00' <= uchar <= '\u9fa5'


def _check_template_modify_time():
    mdate = date.fromtimestamp(osp.getmtime(cfg.excel_path))
    return mdate == cfg.today


def open_excel(path, sheet_name=0):
    try:
        df = pd.read_excel(path, sheet_name, encoding='utf-8')
        logger.info("%s打开成功！" % path)
        return df
    except Exception as e:
        logger.exception(e)
        logger.error("%s打开失败！" % path)
        return None


def run():
    logger.info("开始处理汇总日报……")

    today = cfg.today
    excel_df = open_excel(cfg.excel_path)

    document = Document(osp.join(cfg.proj_path, "template", "Doc1.docx"))

    # 加上标题
    paragraph = document.add_paragraph("xxx团队工作日报（{m}月{d}日）".format(m=today.month, d=today.day))
    paragraph.style = document.styles['m_title']

    # 处理正文
    headers = ["业务1", "业务2", "业务3", "业务4", "业务5"]
    cn_num = "一二三四五六七八九十"
    excel_df_col = list(excel_df.columns)
    for idx, header in enumerate(headers):
        # 新增一段
        header_content = "{idx}、{header}：".format(idx=cn_num[idx], header=header)
        paragraph = document.add_paragraph(header_content)
        paragraph.style = document.styles['m_header']

        # 取出内容
        select_df = excel_df.loc[excel_df[excel_df_col[-1]] == header, excel_df_col[:-1]]
        # 遍历每个人
        for _, row in select_df.iterrows():
            name, items = row[0], row[1]
            if not isinstance(items, str):  # 如果内容为空，则跳过
                continue
            # 遍历每条工作
            for item in items.strip().split('\n'):
                item = item.strip()
                end = len(item) - 1
                # while start < len(item) and not is_chinese(item[start]):
                #     start += 1
                while end >= 0 and item[end] in ";,.；，。":
                    end -= 1
                start = item.find('.') + 1
                item = item[start:end + 1].strip() + "；--" + name

                paragraph = document.add_paragraph(item)
                paragraph.style = document.styles['m_item' + str(idx)]

    save_path = osp.join(cfg.daily_report_path, "%s工作日报.docx" % today.strftime("%m%d"))
    document.save(save_path)
    logger.info("汇总日报处理完了：）")


if __name__ == '__main__':
    if _check_template_modify_time():
        run()
